import os
import tempfile
import logging
from typing import List, Dict, Any, Optional, Union
from datetime import datetime
import hashlib

# LangChain
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain_community.vectorstores import FAISS
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from langchain.schema import Document

# Document Processing
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
import tiktoken

# Google Generative AI
import google.generativeai as genai
from dotenv import load_dotenv

# Utilities
import json

# Load environment variables
load_dotenv()

logger = logging.getLogger(__name__)


class RAGFAISSModel:
    """Modelo RAG usando FAISS y LangChain con Google Generative AI"""

    def __init__(
        self,
        index_path: str = "faiss_index",
        embedding_model: str = "models/embedding-001",
        chat_model: str = "gemini-2.0-flash",
        chunk_size: int = 10000,
        chunk_overlap: int = 1000,
        temperature: float = 0.3,
    ):
        # Configurar Google Generative AI
        api_key = os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY")
        if not api_key:
            raise ValueError(
                "GEMINI_API_KEY o GOOGLE_API_KEY debe estar configurada en el archivo .env"
            )

        genai.configure(api_key=api_key)

        self.index_path = index_path
        self.embedding_model_name = embedding_model
        self.chat_model_name = chat_model
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.temperature = temperature

        # Inicializar componentes
        self.embeddings = GoogleGenerativeAIEmbeddings(
            model=embedding_model, google_api_key=api_key
        )
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size, chunk_overlap=chunk_overlap
        )
        self.tokenizer = tiktoken.get_encoding("cl100k_base")

        # Vector store (se carga cuando existe)
        self.vector_store = None
        self._load_vector_store()

        # Metadatos de documentos
        self.metadata_file = f"{index_path}_metadata.json"
        self.documents_metadata = self._load_metadata()

        logger.info(f"RAG FAISS Model inicializado con embedding: {embedding_model}")

    def _load_vector_store(self):
        """Cargar el vector store si existe"""
        try:
            if os.path.exists(self.index_path):
                self.vector_store = FAISS.load_local(
                    self.index_path,
                    self.embeddings,
                    allow_dangerous_deserialization=True,
                )
                logger.info(f"Vector store cargado desde {self.index_path}")
            else:
                logger.info("No existe vector store previo")
        except Exception as e:
            logger.error(f"Error cargando vector store: {str(e)}")
            self.vector_store = None

    def _load_metadata(self) -> Dict[str, Any]:
        """Cargar metadatos de documentos"""
        try:
            if os.path.exists(self.metadata_file):
                with open(self.metadata_file, "r", encoding="utf-8") as f:
                    return json.load(f)
            return {}
        except Exception as e:
            logger.error(f"Error cargando metadatos: {str(e)}")
            return {}

    def _save_metadata(self):
        """Guardar metadatos de documentos"""
        try:
            with open(self.metadata_file, "w", encoding="utf-8") as f:
                json.dump(self.documents_metadata, f, indent=2, ensure_ascii=False)
        except Exception as e:
            logger.error(f"Error guardando metadatos: {str(e)}")

    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extraer texto de un archivo PDF"""
        text = ""
        try:
            with open(pdf_path, "rb") as file:
                pdf_reader = PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text()
        except Exception as e:
            logger.error(f"Error extrayendo texto de PDF {pdf_path}: {str(e)}")
            raise
        return text

    def extract_text_from_docx(self, docx_path: str) -> str:
        """Extraer texto de un archivo DOCX"""
        text = ""
        try:
            doc = DocxDocument(docx_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        except Exception as e:
            logger.error(f"Error extrayendo texto de DOCX {docx_path}: {str(e)}")
            raise
        return text

    def extract_text_from_txt(self, txt_path: str) -> str:
        """Extraer texto de un archivo de texto"""
        try:
            with open(txt_path, "r", encoding="utf-8") as file:
                return file.read()
        except Exception as e:
            logger.error(f"Error leyendo archivo de texto {txt_path}: {str(e)}")
            raise

    def process_documents(
        self, file_paths: List[str], metadata: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """Procesar múltiples documentos y crear/actualizar el vector store"""
        try:
            all_texts = []
            processed_docs = []

            for file_path in file_paths:
                if not os.path.exists(file_path):
                    logger.warning(f"Archivo no encontrado: {file_path}")
                    continue

                # Extraer texto según el tipo de archivo
                file_ext = os.path.splitext(file_path)[1].lower()

                if file_ext == ".pdf":
                    text = self.extract_text_from_pdf(file_path)
                elif file_ext == ".docx":
                    text = self.extract_text_from_docx(file_path)
                elif file_ext in [".txt", ".md"]:
                    text = self.extract_text_from_txt(file_path)
                else:
                    logger.warning(f"Tipo de archivo no soportado: {file_ext}")
                    continue

                if not text.strip():
                    logger.warning(f"No se pudo extraer texto de: {file_path}")
                    continue

                # Crear chunks del texto
                chunks = self.text_splitter.split_text(text)

                # Generar ID único para el documento
                doc_id = self._generate_document_id(file_path, text)

                # Crear documentos LangChain con metadatos
                for i, chunk in enumerate(chunks):
                    doc_metadata = {
                        "source": file_path,
                        "document_id": doc_id,
                        "chunk_index": i,
                        "total_chunks": len(chunks),
                        "file_type": file_ext,
                        "processed_at": datetime.now().isoformat(),
                        "chunk_tokens": len(self.tokenizer.encode(chunk)),
                    }

                    if metadata:
                        doc_metadata.update(metadata)

                    doc = Document(page_content=chunk, metadata=doc_metadata)
                    processed_docs.append(doc)

                # Guardar información del documento
                self.documents_metadata[doc_id] = {
                    "file_path": file_path,
                    "file_name": os.path.basename(file_path),
                    "file_type": file_ext,
                    "total_chunks": len(chunks),
                    "total_characters": len(text),
                    "total_tokens": sum(
                        len(self.tokenizer.encode(chunk)) for chunk in chunks
                    ),
                    "processed_at": datetime.now().isoformat(),
                    "metadata": metadata or {},
                }

                logger.info(f"Procesado: {file_path} -> {len(chunks)} chunks")

            if not processed_docs:
                raise ValueError("No se pudo procesar ningún documento")

            # Crear o actualizar vector store
            if self.vector_store is None:
                # Crear nuevo vector store
                self.vector_store = FAISS.from_documents(
                    processed_docs, self.embeddings
                )
                logger.info("Nuevo vector store creado")
            else:
                # Agregar documentos al vector store existente
                new_vector_store = FAISS.from_documents(processed_docs, self.embeddings)
                self.vector_store.merge_from(new_vector_store)
                logger.info("Documentos agregados al vector store existente")

            # Guardar vector store
            self.vector_store.save_local(self.index_path)

            # Guardar metadatos
            self._save_metadata()

            result = {
                "status": "success",
                "documents_processed": len(
                    [doc_id for doc_id in self.documents_metadata.keys()]
                ),
                "total_chunks": len(processed_docs),
                "total_tokens": sum(
                    doc["total_tokens"] for doc in self.documents_metadata.values()
                ),
                "processed_at": datetime.now().isoformat(),
            }

            logger.info(f"Procesamiento completado: {result}")
            return result

        except Exception as e:
            logger.error(f"Error procesando documentos: {str(e)}")
            raise

    def search_documents(self, query: str, k: int = 4) -> List[Document]:
        """Buscar documentos similares a la consulta"""
        if self.vector_store is None:
            raise ValueError(
                "No hay vector store disponible. Procesa documentos primero."
            )

        try:
            docs = self.vector_store.similarity_search(query, k=k)
            logger.info(f"Búsqueda realizada: {len(docs)} documentos encontrados")
            return docs
        except Exception as e:
            logger.error(f"Error en búsqueda: {str(e)}")
            raise

    def get_conversational_chain(self) -> Any:
        """Crear cadena conversacional para Q&A"""
        prompt_template = """
        Eres un asistente experto en seguridad química e industrial de BASF.
        Responde en español de manera clara, precisa y profesional.

        Contexto de documentos técnicos:
        {context}

        Pregunta del usuario:
        {question}

        Instrucciones:
        - Proporciona respuestas precisas basadas únicamente en el contexto proporcionado
        - Si la información no está en el contexto, indícalo claramente
        - Incluye datos específicos como números CAS, valores límite, procedimientos de seguridad
        - Estructura la respuesta de manera clara con viñetas cuando sea apropiado
        - Si hay información de seguridad crítica, destácala claramente

        Respuesta:
        """

        # Obtener API key
        api_key = os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY")

        model = ChatGoogleGenerativeAI(
            model=self.chat_model_name,
            temperature=self.temperature,
            google_api_key=api_key,
        )

        prompt = PromptTemplate(
            template=prompt_template, input_variables=["context", "question"]
        )

        chain = load_qa_chain(model, chain_type="stuff", prompt=prompt)
        return chain

    def answer_question(self, question: str, k: int = 4) -> Dict[str, Any]:
        """Responder una pregunta usando RAG"""
        try:
            if self.vector_store is None:
                return {
                    "status": "error",
                    "message": "No hay documentos procesados. Por favor, sube y procesa archivos primero.",
                    "answer": "",
                    "sources": [],
                }

            # Buscar documentos relevantes
            docs = self.search_documents(question, k=k)

            if not docs:
                return {
                    "status": "error",
                    "message": "No se encontraron documentos relevantes para la consulta.",
                    "answer": "",
                    "sources": [],
                }

            # Obtener cadena conversacional
            chain = self.get_conversational_chain()

            # Generar respuesta
            response = chain(
                {"input_documents": docs, "question": question},
                return_only_outputs=True,
            )

            # Extraer fuentes
            sources = []
            for doc in docs:
                source_info = {
                    "source": doc.metadata.get("source", "Desconocido"),
                    "document_id": doc.metadata.get("document_id", ""),
                    "chunk_index": doc.metadata.get("chunk_index", 0),
                    "file_type": doc.metadata.get("file_type", ""),
                }
                if source_info not in sources:
                    sources.append(source_info)

            result = {
                "status": "success",
                "question": question,
                "answer": response["output_text"],
                "sources": sources,
                "context_chunks": len(docs),
                "timestamp": datetime.now().isoformat(),
            }

            logger.info(f"Pregunta respondida exitosamente")
            return result

        except Exception as e:
            logger.error(f"Error respondiendo pregunta: {str(e)}")
            return {
                "status": "error",
                "message": f"Error procesando la consulta: {str(e)}",
                "answer": "",
                "sources": [],
            }

    def get_stats(self) -> Dict[str, Any]:
        """Obtener estadísticas del sistema"""
        try:
            vector_store_exists = self.vector_store is not None

            total_docs = len(self.documents_metadata)
            total_chunks = sum(
                doc["total_chunks"] for doc in self.documents_metadata.values()
            )
            total_tokens = sum(
                doc["total_tokens"] for doc in self.documents_metadata.values()
            )

            return {
                "status": "success",
                "vector_store_exists": vector_store_exists,
                "total_documents": total_docs,
                "total_chunks": total_chunks,
                "total_tokens": total_tokens,
                "embedding_model": self.embedding_model_name,
                "chat_model": self.chat_model_name,
                "index_path": self.index_path,
                "chunk_size": self.chunk_size,
                "chunk_overlap": self.chunk_overlap,
                "documents_metadata": self.documents_metadata,
            }
        except Exception as e:
            logger.error(f"Error obteniendo estadísticas: {str(e)}")
            return {"status": "error", "message": str(e)}

    def _generate_document_id(self, file_path: str, content: str) -> str:
        """Generar un ID único para el documento"""
        content_hash = hashlib.sha256(content.encode()).hexdigest()[:16]
        filename = os.path.basename(file_path)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return f"{filename}_{timestamp}_{content_hash}"


# Instancia global del modelo RAG FAISS
def create_rag_faiss_model():
    """Crear instancia del modelo RAG FAISS"""
    try:
        return RAGFAISSModel()
    except Exception as e:
        logger.error(f"Error creando modelo RAG FAISS: {str(e)}")
        raise


# Instancia global
rag_faiss_model = create_rag_faiss_model()
